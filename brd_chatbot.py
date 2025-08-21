import ollama
import base64
import os
from flask import Flask, request, jsonify, send_file, url_for
import PyPDF2
import tempfile
from pathlib import Path
from doctr.io import DocumentFile
from doctr.models import ocr_predictor
from dotenv import load_dotenv
import logging
import uuid
import torch
from werkzeug.utils import secure_filename
from typing import Dict, List
import shutil
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Preformatted
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
import markdown
import datetime
import graphviz
from PIL import Image, ImageDraw, ImageFont
import io
import re
import json
import hashlib
from flask_cors import CORS
# Try to enable high-fidelity PDF->Word conversion
try:
	from pdf2docx import Converter  # type: ignore
	PDF2DOCX_AVAILABLE = True
except Exception as _e:
	PDF2DOCX_AVAILABLE = False
	logging.warning(f"pdf2docx not available: {_e}")
# Try to enable high-fidelity Word->PDF conversion
try:
    from docx2pdf import convert as docx2pdf_convert  # type: ignore
    DOCX2PDF_AVAILABLE = True
except Exception as _e:
    DOCX2PDF_AVAILABLE = False
    logging.warning(f"docx2pdf not available: {_e}")
import subprocess

# Add minimal DOCX support
from docx import Document

# Optional: MLflow tracking
try:
    import mlflow  # type: ignore
    MLFLOW_AVAILABLE = True
except Exception as _e:
    MLFLOW_AVAILABLE = False
    logging.warning(f"MLflow not available: {_e}")

MLFLOW_TRACKING_URI = os.getenv("MLFLOW_TRACKING_URI")
MLFLOW_EXPERIMENT_NAME = os.getenv("MLFLOW_EXPERIMENT_NAME", "BRDChatbot")

def initialize_mlflow():
    if not MLFLOW_AVAILABLE:
        return
    try:
        if MLFLOW_TRACKING_URI:
            mlflow.set_tracking_uri(MLFLOW_TRACKING_URI)
        mlflow.set_experiment(MLFLOW_EXPERIMENT_NAME)
    except Exception as e:
        logger.warning(f"Failed to initialize MLflow: {str(e)}")

def mlflow_log_event(operation: str, params: Dict = None, metrics: Dict = None, artifacts: List[str] = None, tags: Dict = None):
    if not MLFLOW_AVAILABLE:
        return
    try:
        initialize_mlflow()
        run_name = f"{operation}_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}"
        with mlflow.start_run(run_name=run_name, nested=False):
            if tags:
                try:
                    mlflow.set_tags(tags)
                except Exception:
                    pass
            if params:
                try:
                    mlflow.log_params(params)
                except Exception:
                    pass
            if metrics:
                try:
                    mlflow.log_metrics(metrics)
                except Exception:
                    pass
            if artifacts:
                for path in artifacts:
                    try:
                        if path and os.path.exists(path) and os.path.getsize(path) > 0:
                            mlflow.log_artifact(path, artifact_path="outputs")
                    except Exception:
                        pass
    except Exception as e:
        logger.warning(f"MLflow logging failed for {operation}: {str(e)}")

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Load environment variables
logger.info("Loading environment variables...")
load_dotenv()

# Initialize Flask app
app = Flask(__name__)
CORS(app)
# Initialize global variables
global chat_history, current_brd_id
chat_history = []
current_brd_id = None

# LangMem optional setup
LANGMEM_ENABLED = os.getenv("LANGMEM_ENABLED", "false").lower() == "true"
LANGMEM_BACKEND = os.getenv("LANGMEM_BACKEND", "ollama").lower()  # 'ollama' | 'langmem'
# Use same Ollama model across the app unless overridden
OLLAMA_MODEL = os.getenv("OLLAMA_MODEL", "gemma3:12b")
try:
    from langmem import create_memory_manager
    from langgraph.store.memory import InMemoryStore
    LANGMEM_AVAILABLE = True
except Exception as e:
    LANGMEM_AVAILABLE = False
    logger.warning(f"LangMem not available: {str(e)}")

memory_store = None
memory_manager = None
long_term_memories: List[Dict] = []

def initialize_langmem():
    global memory_store, memory_manager
    # Only initialize the LangMem manager if requested backend is 'langmem'
    if not (LANGMEM_ENABLED and LANGMEM_AVAILABLE and LANGMEM_BACKEND == 'langmem'):
        return
    try:
        memory_store = InMemoryStore()
        default_model = os.getenv("LANGMEM_MODEL", "anthropic:claude-3-5-sonnet-latest")
        memory_manager = create_memory_manager(
            default_model,
            instructions=(
                "Extract and maintain key facts, requirements, decisions, risks, stakeholders, "
                "and user preferences from BRD-related conversations. Consolidate duplicates; "
                "update existing facts when contradicted by newer info."
            ),
            enable_inserts=True,
            enable_updates=True,
            enable_deletes=False,
        )
        logger.info("LangMem initialized (backend=langmem)")
    except Exception as e:
        logger.warning(f"Failed to initialize LangMem: {str(e)}")

def extract_memories_with_ollama(conversation_messages: List[Dict[str, str]]) -> List[str]:
    """Use the same Ollama chat model to extract concise long-term memories from a conversation."""
    try:
        transcript = []
        for msg in conversation_messages:
            role = msg.get("role", "user")
            content = msg.get("content", "")
            transcript.append(f"{role.upper()}: {content}")
        transcript_text = "\n".join(transcript)[-8000:]

        system_prompt = (
            "You are a memory extractor for a BRD assistant. From the conversation transcript, "
            "extract up to 8 concise, standalone memory items useful for future BRD analysis. "
            "Focus on: stakeholders, goals, functional/non-functional requirements, decisions, risks, "
            "constraints, user preferences, document structure notes. Avoid duplicating existing info. "
            "Return a JSON object with an array field 'memories', where each item is a short string."
        )
        user_prompt = f"Transcript:\n\n{transcript_text}\n\nReturn JSON with key 'memories'."
        resp = ollama.chat(
            model=OLLAMA_MODEL,
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt},
            ],
            options={"temperature": 0.2, "num_ctx": 4096}
        )
        content = resp.get("message", {}).get("content", "").strip()
        # Try to locate JSON
        json_text = content
        if "{" in content and "}" in content:
            start = content.find("{")
            end = content.rfind("}") + 1
            json_text = content[start:end]
        data = json.loads(json_text)
        items = data.get("memories", [])
        return [str(x) for x in items if isinstance(x, (str, int, float))][:8]
    except Exception as e:
        logger.warning(f"Ollama memory extraction failed: {str(e)}")
        return []

def update_long_term_memory(conversation_messages: List[Dict[str, str]], brd_id: str = None):
    if not LANGMEM_ENABLED:
        return
    try:
        extracted_items: List[str] = []
        if LANGMEM_BACKEND == 'ollama':
            extracted_items = extract_memories_with_ollama(conversation_messages)
        elif memory_manager:
            extracted = memory_manager.invoke({"messages": conversation_messages})
            for item in extracted or []:
                try:
                    content = getattr(item, "content", item)
                except Exception:
                    content = str(item)
                # Normalize to string
                extracted_items.append(str(content))
        # Store
        for mem_text in extracted_items:
            long_term_memories.append({
                "id": str(uuid.uuid4()),
                "content": mem_text,
                "timestamp": datetime.datetime.now(datetime.UTC).isoformat(),
                "brd_id": brd_id
            })
        if len(long_term_memories) > 200:
            del long_term_memories[:-200]
    except Exception as e:
        logger.warning(f"Long-term memory update failed: {str(e)}")

def get_last_brd_from_history() -> str:
    """Return the most recent uploaded BRD text from chat_history if present."""
    try:
        for msg in reversed(chat_history):
            if msg.get("role") == "user" and isinstance(msg.get("content"), str):
                content = msg["content"]
                if content.startswith("[PDF Content]\n"):
                    return content.split("\n", 1)[1]
    except Exception:
        pass
    return None

def get_brd_hash(content: str) -> str:
    """Generate a unique ID for a BRD based on its content."""
    if not content:
        return None
    sample = content[:1000].encode('utf-8')
    return hashlib.md5(sample).hexdigest()

def summarize_and_store_brd(pdf_text: str):
    """Create a short BRD summary and store it as a memory item for robust recall."""
    try:
        if not pdf_text:
            return
        prompt = (
            "Summarize the BRD into 8-12 bullet points capturing: objectives, scope, key stakeholders, "
            "functional requirements, non-functional requirements, constraints/risks, and major decisions. "
            "Be concise."
        )
        resp = ollama.chat(
            model=OLLAMA_MODEL,
            messages=[
                {"role": "system", "content": "You are a senior BRD analyst."},
                {"role": "user", "content": f"BRD text (truncated):\n{pdf_text[:6000]}\n\n{prompt}"},
            ],
            options={"temperature": 0.2, "num_ctx": 8192}
        )
        summary = resp.get("message", {}).get("content", "").strip()
        if summary:
            brd_id = get_brd_hash(pdf_text)
            long_term_memories.append({
                "id": str(uuid.uuid4()),
                "content": f"[BRD_SUMMARY]\n{summary}",
                "timestamp": datetime.datetime.now(datetime.UTC).isoformat(),
                "brd_id": brd_id
            })
            if len(long_term_memories) > 200:
                del long_term_memories[:-200]
    except Exception as e:
        logger.warning(f"Failed to summarize/store BRD: {str(e)}")

def get_relevant_memories(query: str, limit: int = 5) -> List[str]:
    if not long_term_memories:
        # Fall back to recent BRD summary if available
        return []
    try:
        terms = set(re.findall(r"\w+", query.lower()))
        scored = []
        for m in long_term_memories:
            text = str(m.get("content", ""))
            t_terms = set(re.findall(r"\w+", text.lower()))
            score = len(terms & t_terms)
            if score > 0:
                scored.append((score, text))
        scored.sort(key=lambda x: x[0], reverse=True)
        results = [s[1] for s in scored[:limit]]
        if not results:
            # If nothing matches, include most recent BRD summary if exists
            for m in reversed(long_term_memories):
                txt = str(m.get("content", ""))
                if txt.startswith("[BRD_SUMMARY]"):
                    results.append(txt)
                    break
        return results
    except Exception:
        return []

# Define directories
TEMP_DIR = "temp_files"
PDF_UPLOAD_FOLDER = 'brd_uploads'
PROCESSED_DOCS_FOLDER = 'processed_brds'
OUTPUT_DIR = 'generated_docs'
IMAGE_DIR = 'generated_images'

# Create necessary directories
for folder in [TEMP_DIR, PDF_UPLOAD_FOLDER, PROCESSED_DOCS_FOLDER, OUTPUT_DIR, IMAGE_DIR]:
    os.makedirs(folder, exist_ok=True)

# Initialize doctr model for text extraction
device = torch.device("cuda" if torch.cuda.is_available() else "cpu")
ocr_model = ocr_predictor(pretrained=True).to(device)

def extract_text_from_pdf(pdf_file) -> Dict:
    """Extract text content from PDF file"""
    try:
        # Save PDF to temporary file
        temp_pdf_path = os.path.join(TEMP_DIR, "temp.pdf")
        pdf_file.save(temp_pdf_path)
        
        # Try normal text extraction first
        pdf_reader = PyPDF2.PdfReader(temp_pdf_path)
        text = []
        
        # Extract text from each page
        for page_num in range(len(pdf_reader.pages)):
            page = pdf_reader.pages[page_num]
            page_text = page.extract_text()
            if page_text.strip():
                text.append(f"\n--- Page {page_num + 1} ---\n")
                text.append(page_text)
        
        # If no text found, use OCR
        if not text:
            logger.info("No extractable text found, using OCR...")
            doc = DocumentFile.from_pdf(temp_pdf_path)
            result = ocr_model(doc)
            
            for page_num, page in enumerate(result.pages):
                text.append(f"\n--- Page {page_num + 1} ---\n")
                text.append(page.render())
        
        extracted_text = "\n".join(text)
        return {
            "success": True,
            "text": extracted_text,
            "error": None
        }
        
    except Exception as e:
        logger.error(f"Error processing PDF: {str(e)}")
        return {
            "success": False,
            "text": None,
            "error": str(e)
        }
    
    finally:
        # Cleanup
        if os.path.exists(temp_pdf_path):
            os.remove(temp_pdf_path)

# Minimal DOCX extractor, reusing existing patterns
def extract_text_from_docx(docx_file) -> Dict:
    """Extract text content from Word document (.docx)"""
    try:
        temp_docx_path = os.path.join(TEMP_DIR, "temp.docx")
        docx_file.save(temp_docx_path)

        doc = Document(temp_docx_path)
        text_parts: List[str] = []

        for paragraph in doc.paragraphs:
            if paragraph.text and paragraph.text.strip():
                text_parts.append(paragraph.text)

        for table in doc.tables:
            for row in table.rows:
                row_cells = [cell.text.strip() for cell in row.cells if cell.text and cell.text.strip()]
                if row_cells:
                    text_parts.append(" | ".join(row_cells))

        extracted_text = "\n".join(text_parts)
        return {
            "success": True,
            "text": extracted_text,
            "error": None
        }
    except Exception as e:
        logger.error(f"Error processing Word document: {str(e)}")
        return {
            "success": False,
            "text": None,
            "error": str(e)
        }
    finally:
        try:
            if os.path.exists(temp_docx_path):
                os.remove(temp_docx_path)
        except Exception:
            pass

def clear_chat_history():
    """Clear the chat history"""
    global chat_history
    chat_history = []

def add_to_chat_history(message):
    """Add a message to chat history"""
    global chat_history
    chat_history.append(message)
    # Keep only last 20 messages
    chat_history = chat_history[-20:]

def get_chat_history():
    """Get current chat history"""
    global chat_history
    return chat_history

def get_brd_system_prompt() -> str:
    """Return the system prompt for BRD processing"""
    return (
        "You are an expert Business Requirements Document (BRD) analyst and technical writer. "
        "Your capabilities include:\n"
        "1. Analyzing and understanding BRD content\n"
        "2. Answering questions about the BRD\n"
        "3. Suggesting improvements and updates to the BRD\n"
        "4. Creating architecture diagrams (in text form)\n"
        "5. Creating flowcharts (in text form)\n"
        "6. Extracting key requirements\n"
        "7. Identifying missing information\n\n"
        "When analyzing BRDs:\n"
        "- Focus on clarity and completeness\n"
        "- Identify functional and non-functional requirements\n"
        "- Suggest improvements for unclear sections\n"
        "- Help create visual representations\n"
        "- Maintain professional technical writing standards\n\n"
        "For diagrams and flowcharts:\n"
        "- Use ASCII art or markdown formatting\n"
        "- Keep diagrams clear and readable\n"
        "- Include legends when necessary\n\n"
        "Remember to:\n"
        "- Be specific in your responses\n"
        "- Cite relevant sections of the BRD\n"
        "- Suggest concrete improvements\n"
        "- Maintain document structure and formatting"
    )

def process_brd_query(content: str, pdf_content: str = None) -> str:
    """Process user query about the BRD"""
    try:
        # Create messages for the model
        messages = [
            {
                "role": "system",
                "content": get_brd_system_prompt()
            }
        ]

        # Add BRD content if available or fallback to last BRD from history
        final_brd_context = pdf_content
        if not final_brd_context:
            final_brd_context = get_last_brd_from_history()
        if final_brd_context:
            # Track BRD context id
            global current_brd_id
            current_brd_id = get_brd_hash(final_brd_context)
            messages.append({
                "role": "user",
                "content": f"Here is the BRD content:\n\n{final_brd_context}\n\nPlease keep this content in mind for the next query."
            })

        # Add user query
        messages.append({
            "role": "user",
            "content": content
        })

        # Prepend relevant long-term memory as system context
        recalled = get_relevant_memories(content, limit=5)
        if recalled:
            memory_block = "Relevant long-term memories:\n" + "\n".join(f"- {m}" for m in recalled)
            messages.insert(1, {"role": "system", "content": memory_block})

        # Get response from Gemma model
        response = ollama.chat(
            model=OLLAMA_MODEL,
            messages=messages,
            options={"max_tokens": 2048}
        )

        response_text = response.get("message", {}).get("content", "No response from model")

        # Update long-term memory
        conversation_for_memory: List[Dict[str, str]] = []
        if final_brd_context:
            conversation_for_memory.append({"role": "user", "content": f"BRD context:\n{final_brd_context[:4000]}"})
        conversation_for_memory.append({"role": "user", "content": content})
        conversation_for_memory.append({"role": "assistant", "content": response_text})
        update_long_term_memory(conversation_for_memory, brd_id=current_brd_id)

        return response_text

    except Exception as e:
        logger.error(f"Error processing BRD query: {str(e)}")
        return f"Error processing query: {str(e)}"

def generate_pdf_output(content: str, title: str = "Generated Document") -> Dict:
    """Generate a PDF document from the content with improved formatting"""
    try:
        # Create unique filename
        timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"{title.replace(' ', '_')}_{timestamp}.pdf"
        output_path = os.path.join(OUTPUT_DIR, filename)

        # Create PDF document
        doc = SimpleDocTemplate(
            output_path,
            pagesize=letter,
            rightMargin=72,
            leftMargin=72,
            topMargin=72,
            bottomMargin=72
        )

        # Get styles
        styles = getSampleStyleSheet()
        
        # Create custom styles
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Title'],
            fontSize=24,
            spaceAfter=30,
            alignment=1  # Center alignment
        )

        heading_style = ParagraphStyle(
            'Heading1',
            parent=styles['Heading1'],
            fontSize=16,
            spaceBefore=20,
            spaceAfter=10,
            textColor=colors.HexColor('#2C3E50')
        )

        subheading_style = ParagraphStyle(
            'Heading2',
            parent=styles['Heading2'],
            fontSize=14,
            spaceBefore=15,
            spaceAfter=8,
            textColor=colors.HexColor('#34495E')
        )

        normal_style = ParagraphStyle(
            'CustomNormal',
            parent=styles['Normal'],
            fontSize=11,
            spaceBefore=6,
            spaceAfter=6,
            leading=14
        )

        bullet_style = ParagraphStyle(
            'CustomBullet',
            parent=styles['Normal'],
            fontSize=11,
            spaceBefore=3,
            spaceAfter=3,
            leftIndent=20,
            leading=14
        )

        code_style = ParagraphStyle(
            'CodeStyle',
            parent=styles['Code'],
            fontName='Courier',
            fontSize=9,
            leading=12,
            spaceAfter=10,
            backColor=colors.HexColor('#F7F9FA')
        )

        # Create content elements
        elements = []
        
        # Add title
        elements.append(Paragraph(title, title_style))
        elements.append(Spacer(1, 20))

        # Process content
        current_section = []
        in_code_block = False
        code_content = []

        for line in content.split('\n'):
            line = line.strip()
            if not line:
                # Process accumulated section
                if current_section:
                    section_text = ' '.join(current_section)
                    elements.append(Paragraph(section_text, normal_style))
                    current_section = []
                continue

            if line.startswith('```'):
                # Handle code blocks
                if current_section:
                    elements.append(Paragraph(' '.join(current_section), normal_style))
                    current_section = []
                
                in_code_block = not in_code_block
                if not in_code_block and code_content:
                    elements.append(Preformatted('\n'.join(code_content), code_style))
                    code_content = []
            elif in_code_block:
                code_content.append(line)
            elif line.startswith('**') and line.endswith('**'):
                # Handle headings
                if current_section:
                    elements.append(Paragraph(' '.join(current_section), normal_style))
                    current_section = []
                heading_text = line.strip('*').strip()
                if heading_text.startswith(('I.', 'II.', 'III.', 'IV.', 'V.')):
                    elements.append(Paragraph(heading_text, heading_style))
                else:
                    elements.append(Paragraph(heading_text, subheading_style))
            elif line.startswith('*'):
                # Handle bullet points
                if current_section:
                    elements.append(Paragraph(' '.join(current_section), normal_style))
                    current_section = []
                bullet_text = line.strip('*').strip()
                elements.append(Paragraph(f"• {bullet_text}", bullet_style))
            else:
                current_section.append(line)

        # Add any remaining content
        if current_section:
            elements.append(Paragraph(' '.join(current_section), normal_style))
        if code_content:
            elements.append(Preformatted('\n'.join(code_content), code_style))

        # Build PDF
        doc.build(elements)
        
        return {
            "success": True,
            "path": output_path,
            "filename": filename
        }
    
    except Exception as e:
        logger.error(f"Error generating PDF: {str(e)}")
        return {
            "success": False,
            "error": str(e)
        }

def generate_docx_output(content: str, title: str = "Generated Document") -> str:
    """Generate a Word document from the content"""
    try:
        # Create unique filename
        timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"{title.replace(' ', '_')}_{timestamp}.docx"
        output_path = os.path.join(OUTPUT_DIR, filename)

        # Create document
        doc = Document()
        
        # Add title
        doc.add_heading(title, 0)
        
        # Add content
        current_block = []
        in_code_block = False
        
        for line in content.split('\n'):
            if line.startswith('```'):
                if in_code_block:
                    # End code block
                    if current_block:
                        code_text = '\n'.join(current_block)
                        p = doc.add_paragraph()
                        p.add_run(code_text).font.name = 'Courier New'
                    current_block = []
                    in_code_block = False
                else:
                    # Start code block
                    in_code_block = True
            elif in_code_block:
                current_block.append(line)
            else:
                # Regular text
                if line.strip():
                    doc.add_paragraph(line)
        
        # Save document
        doc.save(output_path)
        return output_path
    
    except Exception as e:
        logger.error(f"Error generating DOCX: {str(e)}")
        return None

def generate_markdown_output(content: str, title: str = "Generated Document") -> str:
    """Generate a Markdown document from the content"""
    try:
        # Create unique filename
        timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"{title.replace(' ', '_')}_{timestamp}.md"
        output_path = os.path.join(OUTPUT_DIR, filename)

        # Create markdown content
        md_content = f"# {title}\n\n{content}"
        
        # Save markdown file
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(md_content)
            
        return output_path
    
    except Exception as e:
        logger.error(f"Error generating Markdown: {str(e)}")
        return None

def generate_diagram_image(content: str) -> Dict:
    """Generate image from diagram description"""
    try:
        # Check if content contains flowchart or diagram markers
        if '```mermaid' in content or '```flowchart' in content or '```diagram' in content:
            # Extract diagram content
            pattern = r'```(?:mermaid|flowchart|diagram)\n(.*?)\n```'
            matches = re.findall(pattern, content, re.DOTALL)
            
            if matches:
                diagram_content = matches[0]
                # Create unique filename
                timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
                filename = f"diagram_{timestamp}.png"
                output_path = os.path.join(IMAGE_DIR, filename)
                
                # Create diagram using graphviz
                dot = graphviz.Digraph(comment='Generated Diagram')
                dot.attr(rankdir='LR')  # Left to right layout
                
                # Parse and add nodes/edges
                lines = diagram_content.strip().split('\n')
                for line in lines:
                    if '->' in line:
                        # Handle connections
                        parts = line.split('->')
                        from_node = parts[0].strip()
                        to_node = parts[1].strip()
                        dot.edge(from_node, to_node)
                    elif '[' in line and ']' in line:
                        # Handle node definitions
                        node_match = re.match(r'(\w+)\s*\[(.*?)\]', line)
                        if node_match:
                            node_id = node_match.group(1)
                            node_label = node_match.group(2)
                            dot.node(node_id, node_label)
                
                # Render the diagram
                dot.render(output_path, format='png', cleanup=True)
                
                return {
                    "success": True,
                    "path": f"{output_path}.png",
                    "filename": f"{filename}.png",
                    "type": "diagram"
                }
        
        # Check if content contains ASCII art
        elif any(marker in content for marker in ['|', '+', '-', '>', '<', '^', 'v']):
            # Create image from ASCII art
            lines = content.split('\n')
            font = ImageFont.load_default()
            # Calculate image size
            line_height = 20
            max_width = max(len(line) * 10 for line in lines)
            height = len(lines) * line_height
            
            # Create image
            image = Image.new('RGB', (max_width, height), color='white')
            draw = ImageDraw.Draw(image)
            
            # Draw ASCII art
            y = 0
            for line in lines:
                draw.text((10, y), line, font=font, fill='black')
                y += line_height
            
            # Save image
            timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"ascii_art_{timestamp}.png"
            output_path = os.path.join(IMAGE_DIR, filename)
            image.save(output_path)
            
            return {
                "success": True,
                "path": output_path,
                "filename": filename,
                "type": "ascii"
            }
    
    except Exception as e:
        logger.error(f"Error generating image: {str(e)}")
        return {
            "success": False,
            "error": str(e)
        }
    
    return {
        "success": False,
        "error": "No diagram or ASCII art found in content"
    }

def contains_visual_content(content: str) -> bool:
    """Check if content contains visual elements that should be rendered as image"""
    return (
        '```mermaid' in content or
        '```flowchart' in content or
        '```diagram' in content or
        any(marker in content for marker in ['|', '+', '-', '>', '<', '^', 'v'])
    )

# High-fidelity PDF to Word conversion using the original PDF file
# Returns output_path to .docx or None on failure
def convert_pdf_filestorage_to_docx(pdf_file, title: str) -> str:
    if not PDF2DOCX_AVAILABLE or pdf_file is None:
        return None
    temp_pdf_path = None
    try:
        timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"{title.replace(' ', '_')}_{timestamp}.docx"
        output_path = os.path.join(OUTPUT_DIR, filename)
        # Save uploaded PDF to a temp file for conversion
        temp_pdf_path = os.path.join(TEMP_DIR, f"{uuid.uuid4()}.pdf")
        # Ensure we start from the beginning of the stream
        try:
            pdf_file.stream.seek(0)
        except Exception:
            pass
        pdf_file.save(temp_pdf_path)
        # Convert using pdf2docx
        conv = Converter(temp_pdf_path)
        conv.convert(output_path)
        conv.close()
        return output_path
    except Exception as e:
        logger.warning(f"Exact PDF->Word conversion failed: {str(e)}")
        return None
    finally:
        try:
            if temp_pdf_path and os.path.exists(temp_pdf_path):
                os.remove(temp_pdf_path)
        except Exception:
            pass

def convert_docx_filestorage_to_pdf(docx_file, title: str) -> str:
    """High-fidelity DOCX to PDF conversion using docx2pdf or LibreOffice; returns output path or None."""
    if docx_file is None:
        return None
    temp_docx_path = None
    try:
        timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"{title.replace(' ', '_')}_{timestamp}.pdf"
        output_path = os.path.join(OUTPUT_DIR, filename)
        temp_docx_path = os.path.join(TEMP_DIR, f"{uuid.uuid4()}.docx")
        try:
            docx_file.stream.seek(0)
        except Exception:
            pass
        docx_file.save(temp_docx_path)
        # Try docx2pdf first (best on Windows/Mac)
        if DOCX2PDF_AVAILABLE:
            try:
                # docx2pdf can write to a specific output file path
                docx2pdf_convert(temp_docx_path, output_path)
                if os.path.exists(output_path):
                    return output_path
            except Exception as e:
                logger.warning(f"docx2pdf conversion failed: {str(e)}")
        # Try LibreOffice (Linux-compatible) if available
        soffice_bin = shutil.which('soffice') or shutil.which('libreoffice')
        if soffice_bin:
            try:
                subprocess.run([
                    soffice_bin,
                    '--headless',
                    '--convert-to', 'pdf',
                    '--outdir', OUTPUT_DIR,
                    temp_docx_path
                ], check=True, stdout=subprocess.PIPE, stderr=subprocess.PIPE)
                produced = os.path.join(
                    OUTPUT_DIR,
                    os.path.splitext(os.path.basename(temp_docx_path))[0] + '.pdf'
                )
                if os.path.exists(produced):
                    # Rename to our target output_path for consistency
                    try:
                        os.replace(produced, output_path)
                    except Exception:
                        output_path = produced
                    return output_path
            except Exception as e:
                logger.warning(f"LibreOffice conversion failed: {str(e)}")
        # Fallback: extract text and generate a clean PDF (not layout-identical)
        try:
            doc = Document(temp_docx_path)
            text_parts: List[str] = []
            for paragraph in doc.paragraphs:
                if paragraph.text and paragraph.text.strip():
                    text_parts.append(paragraph.text)
            for table in doc.tables:
                for row in table.rows:
                    row_cells = [cell.text.strip() for cell in row.cells if cell.text and cell.text.strip()]
                    if row_cells:
                        text_parts.append(" | ".join(row_cells))
            content_text = "\n".join(text_parts)
            result = generate_pdf_output(content_text, title)
            if result.get('success'):
                return result.get('path')
        except Exception as e:
            logger.warning(f"Text-based DOCX->PDF fallback failed: {str(e)}")
        return None
    except Exception as e:
        logger.warning(f"DOCX->PDF conversion failed: {str(e)}")
        return None
    finally:
        try:
            if temp_docx_path and os.path.exists(temp_docx_path):
                os.remove(temp_docx_path)
        except Exception:
            pass

# Intent detection helpers

def user_requested_pdf_to_word(text: str) -> bool:
    try:
        t = (text or "").lower()
        if not t.strip():
            return False
        patterns = [
            r"convert\s+.*pdf\s+(to|into)\s+word",
            r"pdf\s*(to|->)\s*word",
            r"word\s+version\s+of\s+the\s+pdf",
            r"export\s+.*(to|as)\s+word",
            r"generate\s+.*word\s+document",
            r"provide\s+.*word\s+document",
            r"create\s+.*word\s+document",
            r"make\s+.*word\s+document",
            r"\bdocx\b",
            r"word\s+doc(ument)?\b",
        ]
        for p in patterns:
            if re.search(p, t):
                return True
        if ("pdf" in t and "word" in t) and any(v in t for v in ["convert", "export", "generate", "make", "create", "provide", "give", "turn"]):
            return True
        return False
    except Exception:
        return False

def user_requested_word_to_pdf(text: str) -> bool:
    try:
        t = (text or "").lower()
        if not t.strip():
            return False
        patterns = [
            r"convert\s+.*(word|docx|doc)\s+(to|into)\s+pdf",
            r"(word|docx|doc)\s*(to|->)\s*pdf",
            r"pdf\s+version\s+of\s+the\s+(word|docx|document)",
            r"export\s+.*(to|as)\s+pdf",
            r"generate\s+.*pdf\s+document",
            r"create\s+.*pdf\s+document",
            r"make\s+.*pdf\s+document",
        ]
        for p in patterns:
            if re.search(p, t):
                return True
        if ("pdf" in t and ("word" in t or "docx" in t or "document" in t)) and any(v in t for v in ["convert", "export", "generate", "make", "create", "provide", "give", "turn"]):
            return True
        return False
    except Exception:
        return False

def user_requested_add_content(text: str) -> bool:
    """Detect if user wants to add content to existing document"""
    try:
        t = (text or "").lower()
        if not t.strip():
            return False
        patterns = [
            r"add\s+.*to\s+(pdf|document|doc|docx)",
            r"append\s+.*to\s+(pdf|document|doc|docx)",
            r"insert\s+.*to\s+(pdf|document|doc|docx)",
            r"include\s+.*in\s+(pdf|document|doc|docx)",
            r"update\s+.*in\s+(pdf|document|doc|docx)"
        ]
        for p in patterns:
            if re.search(p, t):
                return True
        if "add" in t and any(doc_type in t for doc_type in ["pdf", "document", "doc", "docx"]):
            return True
        return False
    except Exception:
        return False

def append_to_pdf(pdf_file, content: str) -> Dict:
    """Append content to existing PDF while preserving exact original formatting"""
    temp_pdf_path = None
    temp_content_pdf = None
    try:
        # Guard: ensure there is content to add
        if not content or not content.strip():
            return {
                "success": False,
                "error": "No content to add. Please specify what to add."
            }
        
        # Create temporary file for the original PDF
        temp_pdf_path = os.path.join(TEMP_DIR, f"{uuid.uuid4()}.pdf")
        temp_content_pdf = os.path.join(TEMP_DIR, f"{uuid.uuid4()}_content.pdf")
        # Ensure we start from the beginning of the upload stream
        try:
            pdf_file.stream.seek(0)
        except Exception:
            pass
        pdf_file.save(temp_pdf_path)
        # Validate saved file is not empty
        if not os.path.exists(temp_pdf_path) or os.path.getsize(temp_pdf_path) == 0:
            raise ValueError("Uploaded PDF stream was empty or unreadable. Please re-upload the file.")
        
        # Read the original PDF
        reader = PyPDF2.PdfReader(temp_pdf_path)
        writer = PyPDF2.PdfWriter()
        
        # Get formatting from last page
        last_page = reader.pages[-1]
        page_width = float(last_page.mediabox.width)
        page_height = float(last_page.mediabox.height)
        
        # Copy all pages from original PDF with exact formatting
        for page in reader.pages:
            writer.add_page(page)
        
        # Create a new PDF with the additional content matching original format
        timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"updated_document_{timestamp}.pdf"
        output_path = os.path.join(OUTPUT_DIR, filename)
        
        # Generate content PDF with matching page size and margins
        doc = SimpleDocTemplate(
            temp_content_pdf,
            pagesize=(page_width, page_height),
            rightMargin=36,  # Standard margins
            leftMargin=36,
            topMargin=36,
            bottomMargin=36
        )
        
        # Copy styles from original if possible
        styles = getSampleStyleSheet()
        normal_style = ParagraphStyle(
            'OriginalNormal',
            parent=styles['Normal'],
            fontSize=11,  # Default size
            leading=14    # Default leading
        )
        
        # Format new content
        elements = []
        for paragraph in content.split('\n'):
            if paragraph.strip():
                elements.append(Paragraph(paragraph.strip(), normal_style))
        
        # Build content PDF
        doc.build(elements)
        
        # Validate generated content PDF
        if not os.path.exists(temp_content_pdf) or os.path.getsize(temp_content_pdf) == 0:
            return {
                "success": False,
                "error": "Failed to generate content to append."
            }
        
        # Add new content pages
        content_reader = PyPDF2.PdfReader(temp_content_pdf)
        for page in content_reader.pages:
            writer.add_page(page)
        
        # Write the combined PDF
        with open(output_path, 'wb') as output_file:
            writer.write(output_file)
        
        return {
            "success": True,
            "path": output_path,
            "filename": filename
        }
        
    except Exception as e:
        logger.error(f"Error in append_to_pdf: {str(e)}")
        return {
            "success": False,
            "error": str(e)
        }
    finally:
        # Cleanup temporary files
        for temp_file in [temp_pdf_path, temp_content_pdf]:
            try:
                if temp_file and os.path.exists(temp_file):
                    os.remove(temp_file)
            except Exception as cleanup_error:
                logger.warning(f"Failed to cleanup temporary file: {cleanup_error}")
                pass

def append_to_docx(docx_file, content: str) -> Dict:
    """Append content to existing Word document while preserving exact original formatting"""
    try:
        # Guard: ensure there is content to add
        if not content or not content.strip():
            return {
                "success": False,
                "error": "No content to add. Please specify what to add."
            }
        
        # Create temporary file for the original DOCX
        temp_docx_path = os.path.join(TEMP_DIR, f"{uuid.uuid4()}.docx")
        # Ensure we start from the beginning of the upload stream
        try:
            docx_file.stream.seek(0)
        except Exception:
            pass
        docx_file.save(temp_docx_path)
        # Validate saved file is not empty
        if not os.path.exists(temp_docx_path) or os.path.getsize(temp_docx_path) == 0:
            raise ValueError("Uploaded Word stream was empty or unreadable. Please re-upload the file.")
        
        # Open the original document
        doc = Document(temp_docx_path)
        
        # Store original document styles and formatting
        original_styles = {}
        if doc.paragraphs:
            last_para = doc.paragraphs[-1]
            original_styles['font'] = last_para.style.font.name
            original_styles['size'] = last_para.style.font.size
            original_styles['bold'] = last_para.style.font.bold
            original_styles['italic'] = last_para.style.font.italic
            original_styles['alignment'] = last_para.alignment
            original_styles['spacing_before'] = last_para.paragraph_format.space_before
            original_styles['spacing_after'] = last_para.paragraph_format.space_after
            original_styles['line_spacing'] = last_para.paragraph_format.line_spacing
        
        # Add a page break before new content
        doc.add_page_break()
        
        # Add the new content with matched formatting
        for paragraph in content.split('\n'):
            if paragraph.strip():
                p = doc.add_paragraph(paragraph.strip())
                # Apply original formatting
                if original_styles:
                    p.style = last_para.style
                    run = p.runs[0] if p.runs else p.add_run()
                    font = run.font
                    if original_styles.get('font'):
                        font.name = original_styles['font']
                    if original_styles.get('size'):
                        font.size = original_styles['size']
                    if original_styles.get('bold') is not None:
                        font.bold = original_styles['bold']
                    if original_styles.get('italic') is not None:
                        font.italic = original_styles['italic']
                    if original_styles.get('alignment') is not None:
                        p.alignment = original_styles['alignment']
                    if original_styles.get('spacing_before') is not None:
                        p.paragraph_format.space_before = original_styles['spacing_before']
                    if original_styles.get('spacing_after') is not None:
                        p.paragraph_format.space_after = original_styles['spacing_after']
                    if original_styles.get('line_spacing') is not None:
                        p.paragraph_format.line_spacing = original_styles['line_spacing']
        
        # Save the updated document
        timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"updated_document_{timestamp}.docx"
        output_path = os.path.join(OUTPUT_DIR, filename)
        doc.save(output_path)
        
        return {
            "success": True,
            "path": output_path,
            "filename": filename
        }
        
    except Exception as e:
        return {
            "success": False,
            "error": str(e)
        }
    finally:
        try:
            if os.path.exists(temp_docx_path):
                os.remove(temp_docx_path)
        except Exception:
            pass

@app.route('/download/<filename>')
def download_file(filename):
    """Handle file downloads"""
    try:
        # Check if it's an image
        if filename.endswith(('.png', '.jpg', '.jpeg')):
            return send_file(
                os.path.join(IMAGE_DIR, filename),
                mimetype=f'image/{filename.split(".")[-1]}',
                as_attachment=True,
                download_name=filename
            )
        # Otherwise assume it's a PDF
        return send_file(
            os.path.join(OUTPUT_DIR, filename),
            as_attachment=True,
            download_name=filename
        )
    except Exception as e:
        logger.error(f"Error downloading file: {str(e)}")
        return jsonify({'error': 'File not found'}), 404

@app.route('/health', methods=['GET'])
def health():
    return jsonify({"status": "ok", "time": datetime.datetime.now(datetime.UTC).isoformat()})

@app.route('/debug/memories', methods=['GET'])
def debug_memories():
    """Debug endpoint to check current memories (disable in production)"""
    try:
        return jsonify({
            "enabled": LANGMEM_ENABLED,
            "backend": LANGMEM_BACKEND,
            "model": OLLAMA_MODEL,
            "memory_count": len(long_term_memories),
            "memories": [
                {
                    "id": m["id"],
                    "content": m["content"],
                    "timestamp": m["timestamp"],
                    "brd_id": m.get("brd_id")
                }
                for m in sorted(long_term_memories, key=lambda x: x["timestamp"], reverse=True)
            ] if LANGMEM_ENABLED else []
        })
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/chat', methods=['POST'])
def chat():
    """Handle chat requests"""
    try:
        initialize_langmem()
        json_body = request.get_json(silent=True) or {}
        user_content = (request.form.get('content') or json_body.get('content') or '').strip()
        pdf_file = request.files.get('pdf')
        docx_file = request.files.get('docx')
        output_format = (request.form.get('output_format') or json_body.get('output_format') or 'json').lower()
        convert_pdf_to_word_flag = (request.form.get('convert_pdf_to_word') or json_body.get('convert_pdf_to_word') or 'false')
        convert_pdf_to_word = str(convert_pdf_to_word_flag).lower() in ['1', 'true', 'yes', 'on']
        # Optional: request to convert uploaded Word directly to PDF
        convert_word_to_pdf_flag = (request.form.get('convert_word_to_pdf') or json_body.get('convert_word_to_pdf') or 'false')
        convert_word_to_pdf = str(convert_word_to_pdf_flag).lower() in ['1', 'true', 'yes', 'on']
        pdf_content = None
        pdf_uploaded = bool(pdf_file)

        if pdf_file:
            logger.info("Processing uploaded PDF...")
            result = extract_text_from_pdf(pdf_file)
            if not result["success"]:
                logger.error(f"PDF processing error: {result['error']}")
                return jsonify({'error': f"Failed to process PDF: {result['error']}"}), 400
            pdf_content = result["text"]
            logger.info("PDF processed successfully")
            summarize_and_store_brd(pdf_content)
        elif docx_file:
            logger.info("Processing uploaded Word document...")
            result = extract_text_from_docx(docx_file)
            if not result["success"]:
                logger.error(f"Word document processing error: {result['error']}")
                return jsonify({'error': f"Failed to process Word document: {result['error']}"}), 400
            pdf_content = result["text"]
            logger.info("Word document processed successfully")
            summarize_and_store_brd(pdf_content)

        if user_content.lower() in ["hi", "hello", "start", "reset"]:
            clear_chat_history()
            welcome_msg = (
                "Hello! I'm your BRD Assistant. I can help you with:\n\n"
                "📄 BRD Analysis:\n"
                "- Understanding requirements\n"
                "- Suggesting improvements\n"
                "- Creating diagrams and flowcharts\n"
                "- Answering questions about the BRD\n\n"
                "🔍 Features:\n"
                "- PDF processing\n"
                "- Requirement extraction\n"
                "- Architecture visualization\n"
                "- Technical writing assistance\n\n"
                "📝 Output Formats:\n"
                "- PDF document (with download link)\n"
                "- Word document (with download link)\n"
                "- Image format (for diagrams/flowcharts)\n"
                "- JSON response\n\n"
                "Please upload a BRD (PDF) and ask any questions!"
            )
            return jsonify({'response': welcome_msg})

        if not user_content and not pdf_content:
            return jsonify({'error': 'Please provide either a question or upload a PDF'}), 400

        # Check if user wants to add content to existing document
        if user_requested_add_content(user_content):
            # Process the query to get the content to add
            content_to_add = process_brd_query(user_content, pdf_content)
            # Guard: ensure non-empty content to add
            if not content_to_add or not content_to_add.strip():
                return jsonify({'error': 'No content found to add. Please specify the content to append.'}), 400
            
            if pdf_file:
                result = append_to_pdf(pdf_file, content_to_add)
                if result["success"]:
                    download_url = url_for('download_file', filename=result["filename"], _external=True)
                    return jsonify({
                        'response': 'Content added to PDF successfully',
                        'download_url': download_url,
                        'message': 'Updated PDF with new content'
                    })
                else:
                    return jsonify({'error': f"Failed to add content to PDF: {result.get('error', 'Unknown error')}"}), 500
                    
            elif docx_file:
                result = append_to_docx(docx_file, content_to_add)
                if result["success"]:
                    download_url = url_for('download_file', filename=result["filename"], _external=True)
                    return jsonify({
                        'response': 'Content added to Word document successfully',
                        'download_url': download_url,
                        'message': 'Updated Word document with new content'
                    })
                else:
                    return jsonify({'error': f"Failed to add content to Word document: {result.get('error', 'Unknown error')}"}), 500
            else:
                return jsonify({'error': 'Please provide a PDF or Word document to add content to'}), 400

        # PDF -> Word conversion branch (intent or explicit flag)
        if (convert_pdf_to_word or user_requested_pdf_to_word(user_content)) and pdf_uploaded:
            title = f"Converted BRD"
            output_path = convert_pdf_filestorage_to_docx(pdf_file, title)
            if not output_path and pdf_content:
                output_path = generate_docx_output(pdf_content, title)
            if output_path:
                filename = os.path.basename(output_path)
                download_url = url_for('download_file', filename=filename, _external=True)
                # MLflow log artifact
                mlflow_log_event(
                    operation="pdf_to_word",
                    params={
                        "ollama_model": OLLAMA_MODEL,
                        "pdf_uploaded": True,
                        "method": "pdf2docx" if PDF2DOCX_AVAILABLE else "text_fallback",
                    },
                    artifacts=[output_path],
                    tags={"status": "success"}
                )
                return jsonify({
                    'response': 'PDF converted to Word successfully',
                    'download_url': download_url,
                    'message': 'Word document generated from uploaded PDF'
                })
            else:
                mlflow_log_event(
                    operation="pdf_to_word",
                    params={"pdf_uploaded": True},
                    tags={"status": "failed"}
                )
                return jsonify({'error': 'Failed to convert PDF to Word document'}), 500

        # Word -> PDF conversion branch (intent or explicit flag)
        if (convert_word_to_pdf or user_requested_word_to_pdf(user_content)) and docx_file is not None:
            title = f"Converted BRD"
            pdf_output_path = convert_docx_filestorage_to_pdf(docx_file, title)
            if pdf_output_path:
                filename = os.path.basename(pdf_output_path)
                download_url = url_for('download_file', filename=filename, _external=True)
                # MLflow log artifact
                mlflow_log_event(
                    operation="word_to_pdf",
                    params={
                        "ollama_model": OLLAMA_MODEL,
                        "docx_uploaded": True,
                        "docx2pdf_available": DOCX2PDF_AVAILABLE,
                    },
                    artifacts=[pdf_output_path],
                    tags={"status": "success"}
                )
                return jsonify({
                    'response': 'Word converted to PDF successfully',
                    'download_url': download_url,
                    'message': 'PDF generated from uploaded Word document'
                })
            else:
                mlflow_log_event(
                    operation="word_to_pdf",
                    params={"docx_uploaded": True},
                    tags={"status": "failed"}
                )
                return jsonify({'error': 'Failed to convert Word document to PDF'}), 500

        # Process the query
        response = process_brd_query(user_content, pdf_content)
        
        # Add to chat history
        if pdf_content:
            add_to_chat_history({"role": "user", "content": f"[PDF Content]\n{pdf_content}"})
        add_to_chat_history({"role": "user", "content": user_content})
        add_to_chat_history({"role": "assistant", "content": response})

        # Check if response contains visual content
        has_visual = contains_visual_content(response)

        # Generate output based on requested format
        if output_format == 'json':
            # Only return JSON text response; do not include image URLs
            mlflow_log_event(
                operation="chat_json",
                params={
                    "ollama_model": OLLAMA_MODEL,
                    "has_visual": has_visual,
                    "langmem_enabled": LANGMEM_ENABLED,
                    "langmem_backend": LANGMEM_BACKEND,
                },
                metrics={
                    "prompt_len": len(user_content or ""),
                    "response_len": len(response or ""),
                    "pdf_text_len": len(pdf_content or "") if pdf_content else 0,
                },
                tags={"status": "success"}
            )
            return jsonify({'response': response})
        
        elif output_format == 'pdf':
            title = f"BRD Analysis - {datetime.datetime.now().strftime('%Y-%m-%d')}"
            result = generate_pdf_output(response, title)
            
            if result["success"]:
                download_url = url_for('download_file', filename=result["filename"], _external=True)
                # MLflow log artifact
                mlflow_log_event(
                    operation="chat_pdf",
                    params={"ollama_model": OLLAMA_MODEL},
                    artifacts=[result["path"]],
                    tags={"status": "success"}
                )
                return jsonify({
                    'response': response,
                    'download_url': download_url,
                    'message': 'PDF generated successfully'
                })
            else:
                mlflow_log_event(
                    operation="chat_pdf",
                    params={"ollama_model": OLLAMA_MODEL},
                    tags={"status": "failed", "error": str(result.get('error'))}
                )
                return jsonify({'error': f"Failed to generate PDF: {result.get('error', 'Unknown error')}"}), 500
        
        elif output_format == 'docx' or output_format == 'word':
            title = f"BRD Analysis - {datetime.datetime.now().strftime('%Y-%m-%d')}"
            output_path = generate_docx_output(response, title)
            
            if output_path:
                filename = os.path.basename(output_path)
                download_url = url_for('download_file', filename=filename, _external=True)
                mlflow_log_event(
                    operation="chat_docx",
                    params={"ollama_model": OLLAMA_MODEL},
                    artifacts=[output_path],
                    tags={"status": "success"}
                )
                return jsonify({
                    'response': response,
                    'download_url': download_url,
                    'message': 'Word document generated successfully'
                })
            else:
                mlflow_log_event(
                    operation="chat_docx",
                    params={"ollama_model": OLLAMA_MODEL},
                    tags={"status": "failed"}
                )
                return jsonify({'error': 'Failed to generate Word document'}), 500
        
        elif output_format == 'image' and has_visual:
            image_result = generate_diagram_image(response)
            if image_result["success"]:
                download_url = url_for('download_file', 
                                     filename=image_result["filename"], 
                                     _external=True)
                mlflow_log_event(
                    operation="chat_image",
                    params={"ollama_model": OLLAMA_MODEL, "type": image_result.get('type')},
                    artifacts=[image_result.get("path")],
                    tags={"status": "success"}
                )
                return jsonify({
                    'response': response,
                    'image_url': download_url,
                    'message': f"Image generated successfully ({image_result['type']})"
                })
            else:
                mlflow_log_event(
                    operation="chat_image",
                    params={"ollama_model": OLLAMA_MODEL},
                    tags={"status": "failed", "error": str(image_result.get('error'))}
                )
                return jsonify({'error': f"Failed to generate image: {image_result.get('error', 'Unknown error')}"}), 500
        else:
            return jsonify({'error': f'Unsupported output format: {output_format}'}), 400

    except Exception as e:
        logger.error(f"Error in chat endpoint: {str(e)}")
        return jsonify({'error': str(e)}), 500

if __name__ == '__main__':
    app.run(debug=True, port=9091, host='0.0.0.0')
